"""
Motor de Plantillas de Documento
=================================

Aplica branding corporativo a documentos Word:
- Portadas personalizadas
- Headers y footers con marca
- Tablas de control de versiones
- Estilos consistentes
"""

from pathlib import Path
from typing import Optional, Dict, Any
from datetime import datetime
import logging

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .converter import BrandConfig, DocumentConfig

logger = logging.getLogger(__name__)


class TemplateEngine:
    """
    Motor de plantillas que aplica branding corporativo
    a documentos Word de forma dinámica.
    """
    
    def __init__(
        self,
        brand_config: BrandConfig,
        doc_config: DocumentConfig
    ):
        """
        Inicializa el motor de plantillas.
        
        Args:
            brand_config: Configuración de marca
            doc_config: Configuración del documento
        """
        self.brand_config = brand_config
        self.doc_config = doc_config
        self.language = doc_config.metadata.get('language', 'es')
    
    def get_text(self, key: str) -> str:
        """
        Obtiene texto localizado según el idioma del documento.
        
        Args:
            key: Clave del texto
        
        Returns:
            Texto localizado
        """
        default_texts = self.brand_config.brand.get('default_texts', {})
        lang_texts = default_texts.get(self.language, {})
        return lang_texts.get(key, key)
    
    def create_cover_page(self, doc: Document) -> None:
        """
        Crea una portada profesional con branding corporativo.
        
        Args:
            doc: Documento de python-docx
        """
        logger.info("Creando portada con branding corporativo...")
        
        # Asegurar que el documento tiene configuración de sección
        if len(doc.sections) == 0:
            doc.add_section()
        
        # Establecer dimensiones de página si no están definidas
        section = doc.sections[0]
        from docx.shared import Cm
        if section.page_width is None:
            section.page_width = Cm(21.0)  # A4 width
        if section.page_height is None:
            section.page_height = Cm(29.7)  # A4 height
        if section.left_margin is None:
            section.left_margin = Cm(2.54)
        if section.right_margin is None:
            section.right_margin = Cm(2.54)
        if section.top_margin is None:
            section.top_margin = Cm(2.54)
        if section.bottom_margin is None:
            section.bottom_margin = Cm(2.54)
        
        # Configuración de la portada
        cover_config = self.brand_config.brand.get('word_styles', {}).get('cover', {})
        
        # Logo (si existe)
        logo_path = self.brand_config.assets.get('logo_vertical') or self.brand_config.assets.get('logo')
        if logo_path:
            logo_full_path = Path(__file__).parent.parent / logo_path
            if logo_full_path.exists():
                # Agregar logo
                logo_paragraph = doc.add_paragraph()
                logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = logo_paragraph.add_run()
                try:
                    run.add_picture(str(logo_full_path), width=Inches(2.5))
                except Exception as e:
                    logger.warning(f"No se pudo agregar logo: {e}")
        
        # Espaciado
        doc.add_paragraph()
        
        # Título principal
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(self.doc_config.title)
        title_run.font.name = 'Montserrat'
        title_run.font.size = Pt(48)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(*self.brand_config.get_color_rgb('weld_orange'))
        
        # Subtítulo (si existe)
        if self.doc_config.subtitle:
            subtitle = doc.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle.add_run(self.doc_config.subtitle)
            subtitle_run.font.name = 'Montserrat'
            subtitle_run.font.size = Pt(24)
            subtitle_run.font.color.rgb = RGBColor(*self.brand_config.get_color_rgb('precision_blue'))
        
        # Espaciado
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Información del documento (en tabla)
        info_table = doc.add_table(rows=0, cols=2)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Configurar ancho de columnas
        info_table.autofit = False
        info_table.allow_autofit = False
        
        # Datos a mostrar
        info_data = []
        
        if self.doc_config.project:
            info_data.append((self.get_text('project'), self.doc_config.project))
        
        if self.doc_config.client:
            info_data.append((self.get_text('client'), self.doc_config.client))
        
        info_data.append((self.get_text('author'), self.doc_config.author))
        info_data.append((self.get_text('version'), self.doc_config.version))
        info_data.append((self.get_text('revision'), self.doc_config.revision))
        info_data.append((self.get_text('date'), self.doc_config.date))
        
        # Llenar tabla
        for label, value in info_data:
            row = info_table.add_row()
            
            # Celda de etiqueta
            label_cell = row.cells[0]
            label_para = label_cell.paragraphs[0]
            label_run = label_para.add_run(label + ":")
            label_run.font.bold = True
            label_run.font.size = Pt(11)
            label_run.font.color.rgb = RGBColor(*self.brand_config.get_color_rgb('steel_gray'))
            
            # Celda de valor
            value_cell = row.cells[1]
            value_para = value_cell.paragraphs[0]
            value_run = value_para.add_run(value)
            value_run.font.size = Pt(11)
            value_run.font.color.rgb = RGBColor(*self.brand_config.get_color_rgb('dark_steel'))
        
        # Espaciado final
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Nota de confidencialidad
        if self.doc_config.confidentiality:
            conf_para = doc.add_paragraph()
            conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            conf_run = conf_para.add_run(self.doc_config.confidentiality)
            conf_run.font.name = 'Inter'
            conf_run.font.size = Pt(10)
            conf_run.font.italic = True
            conf_run.font.color.rgb = RGBColor(*self.brand_config.get_color_rgb('steel_gray'))
        
        # Salto de página
        doc.add_page_break()
        
        logger.info("✓ Portada creada exitosamente")
    
    def create_version_control_table(self, doc: Document) -> None:
        """
        Crea tabla de control de versiones.
        
        Args:
            doc: Documento de python-docx
        """
        if not self.doc_config.version_control:
            return
        
        logger.info("Creando tabla de control de versiones...")
        
        # Título de la sección
        title = doc.add_paragraph()
        title_run = title.add_run(self.get_text('version_control'))
        title_run.font.name = 'Montserrat'
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(*self.brand_config.get_color_rgb('precision_blue'))
        
        doc.add_paragraph()
        
        # Crear tabla
        table = doc.add_table(rows=1, cols=4)
        # Nota: El estilo puede no estar disponible en documentos generados por pandoc
        try:
            table.style = 'Light Grid Accent 1'
        except KeyError:
            # Usar estilo por defecto si no existe
            pass
        
        # Headers
        headers = [
            self.get_text('version'),
            self.get_text('date'),
            self.get_text('author'),
            'Descripción' if self.language == 'es' else 'Description'
        ]
        
        header_cells = table.rows[0].cells
        header_rgb = self.brand_config.get_color_rgb('weld_orange')
        
        for i, header_text in enumerate(headers):
            cell = header_cells[i]
            
            # Texto del header
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(header_text)
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(255, 255, 255)
            
            # Color de fondo
            self._set_cell_background(cell, header_rgb)
        
        # Filas de versiones
        for version_info in self.doc_config.version_control:
            row = table.add_row()
            
            row.cells[0].text = version_info.get('version', '')
            row.cells[1].text = version_info.get('date', '')
            row.cells[2].text = version_info.get('author', '')
            row.cells[3].text = version_info.get('description', '')
            
            # Estilo de las celdas
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)
        
        # Salto de página
        doc.add_page_break()
        
        logger.info("✓ Tabla de control de versiones creada")
    
    def _set_cell_background(self, cell, rgb: tuple):
        """Establece color de fondo de celda"""
        cell_properties = cell._element.get_or_add_tcPr()
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '%02x%02x%02x' % rgb)
        cell_properties.append(shading_elm)
    
    def create_header(self, section) -> None:
        """
        Crea encabezado de página con branding.
        
        Args:
            section: Sección del documento
        """
        logger.info("Creando encabezado...")
        
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Logo pequeño (si existe)
        logo_path = self.brand_config.assets.get('logo_small') or self.brand_config.assets.get('logo')
        if logo_path:
            logo_full_path = Path(__file__).parent.parent / logo_path
            if logo_full_path.exists():
                try:
                    run = header_para.add_run()
                    run.add_picture(str(logo_full_path), height=Inches(0.4))
                except Exception as e:
                    logger.warning(f"No se pudo agregar logo al header: {e}")
            else:
                # Texto alternativo
                run = header_para.add_run(self.brand_config.brand.get('name', ''))
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(*self.brand_config.get_color_rgb('steel_gray'))
        else:
            # Texto alternativo
            run = header_para.add_run(self.brand_config.brand.get('name', ''))
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(*self.brand_config.get_color_rgb('steel_gray'))
        
        # Línea separadora
        header_para = header.add_paragraph()
        header_para.add_run('_' * 80)
        
        logger.info("✓ Encabezado creado")
    
    def create_footer(self, section) -> None:
        """
        Crea pie de página con numeración y confidencialidad.
        
        Args:
            section: Sección del documento
        """
        logger.info("Creando pie de página...")
        
        footer = section.footer
        
        # Tabla para organizar footer (3 columnas)
        footer_table = footer.add_table(rows=1, cols=3, width=Inches(6))
        footer_table.autofit = False
        
        # Columna izquierda: Confidencialidad
        left_cell = footer_table.rows[0].cells[0]
        left_para = left_cell.paragraphs[0]
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_run = left_para.add_run(self.doc_config.confidentiality)
        left_run.font.size = Pt(8)
        left_run.font.color.rgb = RGBColor(*self.brand_config.get_color_rgb('steel_gray'))
        
        # Columna central: Número de página
        center_cell = footer_table.rows[0].cells[1]
        center_para = center_cell.paragraphs[0]
        center_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Agregar campo de número de página
        center_run = center_para.add_run()
        center_run.font.size = Pt(9)
        center_run.font.color.rgb = RGBColor(*self.brand_config.get_color_rgb('steel_gray'))
        
        # Insertar campo de página actual
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        center_run._r.append(fldChar1)
        center_run._r.append(instrText)
        center_run._r.append(fldChar2)
        
        # Columna derecha: Fecha
        right_cell = footer_table.rows[0].cells[2]
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_run = right_para.add_run(self.doc_config.date)
        right_run.font.size = Pt(8)
        right_run.font.color.rgb = RGBColor(*self.brand_config.get_color_rgb('steel_gray'))
        
        # Eliminar bordes de la tabla
        self._remove_table_borders(footer_table)
        
        logger.info("✓ Pie de página creado")
    
    def _remove_table_borders(self, table):
        """Elimina los bordes de una tabla"""
        tbl = table._tbl
        tblPr = tbl.tblPr
        
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        # Crear elemento de bordes
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            tblBorders.append(border)
        
        tblPr.append(tblBorders)
    
    def apply_typography(self, doc: Document) -> None:
        """
        Aplica tipografía y estilos según configuración de marca.
        
        Args:
            doc: Documento de python-docx
        """
        logger.info("Aplicando tipografía corporativa...")
        
        typography = self.brand_config.typography
        
        for paragraph in doc.paragraphs:
            # Determinar tipo de párrafo y aplicar estilo
            style_name = paragraph.style.name
            
            if style_name.startswith('Heading'):
                # Extraer nivel de heading
                level = style_name.replace('Heading ', '').replace('heading ', '')
                style_key = f'h{level}'
                
                if style_key in typography:
                    self._apply_paragraph_typography(paragraph, typography[style_key])
            
            elif style_name == 'Normal':
                # Aplicar estilo de cuerpo
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                if 'body' in typography:
                    self._apply_paragraph_typography(paragraph, typography['body'])
        
        logger.info("✓ Tipografía aplicada")
    
    def _apply_paragraph_typography(self, paragraph, typo_config: Dict):
        """Aplica configuración tipográfica a un párrafo"""
        for run in paragraph.runs:
            # Tamaño
            if 'size' in typo_config:
                run.font.size = Pt(typo_config['size'])
            
            # Peso (negrita)
            if 'weight' in typo_config:
                weight = typo_config['weight']
                if weight >= 700:
                    run.font.bold = True
                elif weight >= 600:
                    run.font.bold = True
            
            # Color
            if 'color' in typo_config:
                color_name = typo_config['color']
                rgb = self.brand_config.get_color_rgb(color_name)
                run.font.color.rgb = RGBColor(*rgb)
            
            # Familia de fuente
            if 'font_family' in typo_config:
                run.font.name = typo_config['font_family']
    
    def format_tables(self, doc: Document) -> None:
        """
        Aplica formato corporativo a todas las tablas del documento.
        
        Args:
            doc: Documento de python-docx
        """
        logger.info("Formateando tablas...")
        
        table_config = self.brand_config.brand.get('word_styles', {}).get('tables', {})
        header_rgb = self.brand_config.get_color_rgb('weld_orange')
        alt_rgb = self.brand_config.get_color_rgb('light_steel')
        
        for table in doc.tables:
            # Aplicar bordes
            self._set_table_borders(table)
            
            # Formatear header (primera fila)
            if table.rows:
                header_row = table.rows[0]
                for cell in header_row.cells:
                    # Color de fondo
                    self._set_cell_background(cell, header_rgb)
                    
                    # Formato de texto
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.font.size = Pt(11)
                
                # Alternar color de filas
                for i, row in enumerate(table.rows[1:], start=1):
                    if i % 2 == 0:
                        for cell in row.cells:
                            self._set_cell_background(cell, alt_rgb)
        
        logger.info("✓ Tablas formateadas")
    
    def _set_table_borders(self, table):
        """Establece bordes para la tabla"""
        tbl = table._tbl
        tblPr = tbl.tblPr
        
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        # Crear elemento de bordes
        tblBorders = OxmlElement('w:tblBorders')
        border_color = self.brand_config.get_color_rgb('steel_gray')
        border_hex = '%02x%02x%02x' % border_color
        
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), border_hex)
            tblBorders.append(border)
        
        tblPr.append(tblBorders)
    
    def apply_full_template(self, doc: Document) -> Document:
        """
        Aplica la plantilla completa al documento:
        - Portada
        - Tabla de control de versiones
        - Headers y footers
        - Tipografía
        - Formato de tablas
        
        Args:
            doc: Documento de python-docx
        
        Returns:
            Documento modificado
        """
        logger.info("Aplicando plantilla completa...")
        
        # Crear portada
        self.create_cover_page(doc)
        
        # Crear tabla de control de versiones
        self.create_version_control_table(doc)
        
        # Aplicar headers y footers a todas las secciones
        for section in doc.sections:
            self.create_header(section)
            self.create_footer(section)
        
        # Aplicar tipografía
        self.apply_typography(doc)
        
        # Formatear tablas
        self.format_tables(doc)
        
        logger.info("✓ Plantilla completa aplicada exitosamente")
        
        return doc

