"""
Motor de Conversión de Markdown a Word
========================================

Clase base para convertir documentos Markdown a Word (.docx) 
con formato profesional y branding configurable.
"""

import os
import re
from pathlib import Path
from datetime import datetime
from typing import Optional, Dict, Any, List
import logging

try:
    import pypandoc
except ImportError:
    pypandoc = None

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    Document = None

# Configurar logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class BrandConfig:
    """Configuración de marca para aplicar al documento"""
    def __init__(self, config_dict: Dict[str, Any]):
        self.brand = config_dict.get('brand', {})
        self.colors = config_dict.get('colors', {})
        self.fonts = config_dict.get('fonts', {})
        self.typography = config_dict.get('typography', {})
        self.assets = config_dict.get('assets', {})
    
    def get_color_rgb(self, color_name: str) -> tuple:
        """Convierte color hex a RGB tuple"""
        # Buscar el color en la configuración
        for category in self.colors.values():
            if isinstance(category, dict) and color_name in category:
                hex_color = category[color_name]
                return self._hex_to_rgb(hex_color)
        
        # Si no se encuentra, retornar negro por defecto
        return (0, 0, 0)
    
    def _hex_to_rgb(self, hex_color: str) -> tuple:
        """Convierte hex a RGB"""
        hex_color = hex_color.lstrip('#')
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


class DocumentConfig:
    """Configuración del documento a generar"""
    def __init__(self, config_dict: Dict[str, Any]):
        self.document = config_dict.get('document', {})
        self.version_control = config_dict.get('version_control', [])
        self.metadata = config_dict.get('metadata', {})
    
    @property
    def title(self) -> str:
        return self.document.get('title', 'Documento Sin Título')
    
    @property
    def subtitle(self) -> str:
        return self.document.get('subtitle', '')
    
    @property
    def project(self) -> str:
        return self.document.get('project', '')
    
    @property
    def client(self) -> str:
        return self.document.get('client', '')
    
    @property
    def author(self) -> str:
        return self.document.get('author', '')
    
    @property
    def version(self) -> str:
        return self.document.get('version', '1.0')
    
    @property
    def revision(self) -> str:
        return self.document.get('revision', 'RA-00')
    
    @property
    def date(self) -> str:
        return self.document.get('date', datetime.now().strftime('%Y-%m-%d'))
    
    @property
    def confidentiality(self) -> str:
        return self.document.get('confidentiality', 'DOCUMENTO TÉCNICO')


class DocumentConverter:
    """
    Clase base para conversión de documentos Markdown a Word
    con soporte para branding configurable.
    """
    
    def __init__(
        self, 
        input_file: Path,
        output_file: Optional[Path] = None,
        brand_config: Optional[BrandConfig] = None,
        doc_config: Optional[DocumentConfig] = None
    ):
        self.input_file = Path(input_file)
        self.output_file = Path(output_file) if output_file else self.input_file.with_suffix('.docx')
        self.brand_config = brand_config
        self.doc_config = doc_config
        self.doc = None
        
        # Verificar dependencias
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Verifica que las dependencias estén instaladas"""
        if pypandoc is None:
            raise ImportError(
                "pypandoc no está instalado. Instálalo con: pip install pypandoc"
            )
        if Document is None:
            raise ImportError(
                "python-docx no está instalado. Instálalo con: pip install python-docx"
            )
    
    def verify_pandoc(self) -> bool:
        """Verifica que pandoc esté instalado"""
        try:
            version = pypandoc.get_pandoc_version()
            logger.info(f"✓ Pandoc versión: {version}")
            return True
        except Exception as e:
            logger.warning(f"⚠️  Pandoc no encontrado: {e}")
            try:
                pypandoc.download_pandoc()
                logger.info("✓ Pandoc descargado correctamente")
                return True
            except Exception as e:
                logger.error(f"✗ Error al instalar pandoc: {e}")
                logger.info("\nInstala pandoc manualmente:")
                logger.info("  Windows: winget install --id JohnMacFarlane.Pandoc")
                logger.info("  Linux: sudo apt install pandoc")
                logger.info("  Mac: brew install pandoc")
                return False
    
    def convert_markdown_basic(self) -> Optional[Path]:
        """Convierte el archivo Markdown a Word usando pypandoc"""
        logger.info(f"\n📄 Convirtiendo: {self.input_file}")
        
        temp_file = self.output_file.with_name(
            self.output_file.stem + '_temp' + self.output_file.suffix
        )
        
        try:
            # Opciones de conversión
            extra_args = [
                '--toc',                           # Tabla de contenido
                '--toc-depth=3',                   # Profundidad hasta nivel 3
            ]
            
            # Agregar numeración de secciones si está configurado
            if self.doc_config and self.doc_config.metadata.get('number_sections', True):
                extra_args.append('--number-sections')
            
            # Conversión con pypandoc
            pypandoc.convert_file(
                str(self.input_file),
                'docx',
                outputfile=str(temp_file),
                extra_args=extra_args
            )
            
            logger.info(f"✓ Conversión básica completada: {temp_file}")
            return temp_file
            
        except Exception as e:
            logger.error(f"✗ Error en conversión: {e}")
            return None
    
    def apply_formatting(self, doc: Document) -> Document:
        """
        Aplica formato profesional al documento.
        Debe ser implementado por subclases o el template engine.
        """
        # Aplicar estilos básicos
        if self.brand_config:
            self._apply_typography(doc)
            self._format_tables(doc)
        
        return doc
    
    def _apply_typography(self, doc: Document):
        """Aplica tipografía según configuración de marca"""
        if not self.brand_config or not self.brand_config.typography:
            return
        
        typography = self.brand_config.typography
        
        for paragraph in doc.paragraphs:
            # Aplicar estilos según el nivel
            if paragraph.style.name.startswith('Heading'):
                level = paragraph.style.name.replace('Heading ', '')
                style_key = f'h{level}'
                
                if style_key in typography:
                    style_config = typography[style_key]
                    self._apply_paragraph_style(paragraph, style_config)
            
            # Justificar párrafos de texto normal
            elif paragraph.style.name == 'Normal':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                if 'body' in typography:
                    self._apply_paragraph_style(paragraph, typography['body'])
    
    def _apply_paragraph_style(self, paragraph, style_config: Dict):
        """Aplica estilo a un párrafo"""
        for run in paragraph.runs:
            if 'size' in style_config:
                run.font.size = Pt(style_config['size'])
            
            if 'weight' in style_config:
                weight = style_config['weight']
                if weight >= 700:
                    run.font.bold = True
            
            if 'color' in style_config and self.brand_config:
                color_name = style_config['color']
                rgb = self.brand_config.get_color_rgb(color_name)
                run.font.color.rgb = RGBColor(*rgb)
    
    def _format_tables(self, doc: Document):
        """Aplica formato profesional a tablas"""
        if not self.brand_config:
            return
        
        # Obtener color primario para headers
        primary_rgb = self.brand_config.get_color_rgb('weld_orange')
        
        for table in doc.tables:
            # Aplicar bordes
            self._set_table_borders(table)
            
            # Formatear encabezado (primera fila)
            if table.rows:
                header_row = table.rows[0]
                for cell in header_row.cells:
                    # Color de fondo
                    self._set_cell_background(cell, primary_rgb)
                    
                    # Texto blanco y en negrita
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.font.size = Pt(11)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _set_table_borders(self, table):
        """Establece bordes para la tabla"""
        tbl = table._tbl
        tblPr = tbl.tblPr
        
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        # Crear elemento de bordes
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        
        tblPr.append(tblBorders)
    
    def _set_cell_background(self, cell, rgb: tuple):
        """Establece color de fondo de celda"""
        cell_properties = cell._element.get_or_add_tcPr()
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '%02x%02x%02x' % rgb)
        cell_properties.append(shading_elm)
    
    def convert(self) -> bool:
        """
        Proceso completo de conversión.
        Retorna True si fue exitoso.
        """
        try:
            # Verificar pandoc
            if not self.verify_pandoc():
                return False
            
            # Convertir Markdown a Word
            temp_file = self.convert_markdown_basic()
            if not temp_file or not temp_file.exists():
                logger.error("Error: No se pudo crear el archivo temporal")
                return False
            
            # Abrir documento temporal
            self.doc = Document(str(temp_file))
            
            # Aplicar formato
            self.doc = self.apply_formatting(self.doc)
            
            # Guardar documento final
            self.doc.save(str(self.output_file))
            
            # Eliminar archivo temporal
            if temp_file.exists():
                temp_file.unlink()
            
            logger.info(f"✓ Conversión completada: {self.output_file}")
            return True
            
        except Exception as e:
            logger.error(f"✗ Error durante la conversión: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def create_cover(self):
        """
        Crea página de portada.
        Implementado por TemplateEngine.
        """
        pass
    
    def create_header(self):
        """
        Crea encabezado de página.
        Implementado por TemplateEngine.
        """
        pass
    
    def create_footer(self):
        """
        Crea pie de página.
        Implementado por TemplateEngine.
        """
        pass

